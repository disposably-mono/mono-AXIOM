"""
End-to-end demo of axiom-fetch.

Run the store server in one terminal:
    python -m axiom_store.server --vault /home/mono/Projects/mono-axiom/mono-vault -v

Then run this script in another terminal:
    python scripts/demo_axiom_fetch.py

The script starts a tiny local HTTP server that serves generated HTML,
plain text, PDF, and DOCX documents. Each URL is ingested through the real
fetch pipeline and written to the vault via axiom-store. The generated
source and chunk paths are printed for inspection.
"""

from __future__ import annotations

import argparse
import contextlib
import os
import sys
import threading
from dataclasses import dataclass
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path

from axiom_fetch.pipeline import CHUNKS_DIR, SOURCES_DIR, ingest
from axiom_store import StoreClient, StoreError, parse_frontmatter
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
VAULT_ROOT = REPO_ROOT / "mono-vault"

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass(frozen=True)
class DemoDocument:
    path: str
    content_type: str
    body: bytes
    expected_title: str | None
    expected_text: str


def main() -> int:
    parser = argparse.ArgumentParser(description="Run an end-to-end axiom-fetch demo.")
    parser.add_argument("--host", default="127.0.0.1", help="axiom-store host")
    parser.add_argument(
        "--port",
        type=int,
        default=int(os.environ.get("STORE_PORT", "7070")),
        help="axiom-store port (default: $STORE_PORT or 7070)",
    )
    parser.add_argument(
        "--vault",
        type=Path,
        default=Path(os.environ.get("VAULT_PATH", VAULT_ROOT)),
        help="Vault path used for printed inspection paths.",
    )
    args = parser.parse_args()

    vault_root = args.vault.expanduser().resolve()
    store = StoreClient(host=args.host, port=args.port)
    documents = _demo_documents()

    try:
        _assert_store_is_running(store)
    except OSError as exc:
        print("Could not connect to axiom-store.", file=sys.stderr)
        print("Start it with:", file=sys.stderr)
        print(
            f"  python -m axiom_store.server --vault {vault_root} "
            f"--host {args.host} --port {args.port} -v",
            file=sys.stderr,
        )
        print(f"Connection error: {exc}", file=sys.stderr)
        return 1

    with _serve_documents(documents) as base_url:
        print(f"demo HTTP server: {base_url}")
        print()

        for document in documents:
            url = f"{base_url}{document.path}"
            print(f"=== ingest {document.path} ===")
            source_id = ingest(url, store)
            source_path = f"{SOURCES_DIR}{source_id}.md"

            source_meta, _ = parse_frontmatter(store.read(source_path).decode("utf-8"))
            assert source_meta["status"] == "succeeded", source_meta
            assert source_meta["content_type"] == document.content_type
            assert source_meta["chunk_count"] >= 1
            if document.expected_title is not None:
                assert source_meta["title"] == document.expected_title

            chunk_path = f"{CHUNKS_DIR}{source_id}-0000.md"
            chunk_meta, chunk_body = parse_frontmatter(store.read(chunk_path).decode("utf-8"))
            assert chunk_meta["source_id"] == source_id
            assert document.expected_text in chunk_body

            print(f"  source: {vault_root / source_path}")
            print(f"  first chunk: {vault_root / chunk_path}")
            print(f"  title: {source_meta.get('title', '(none)')}")
            print(f"  chunks: {source_meta['chunk_count']}")
            print()

    print("All axiom-fetch demo ingests passed.")
    return 0


def _assert_store_is_running(store: StoreClient) -> None:
    """Perform a cheap read to fail fast with a friendly message."""
    try:
        store.list_dir("fetch/sources")
    except FileNotFoundError:
        # The vault may not have the fetch directory yet; any server response
        # proves the store is reachable.
        return
    except StoreError:
        raise


@contextlib.contextmanager
def _serve_documents(documents: list[DemoDocument]):
    routes = {document.path: document for document in documents}

    class Handler(BaseHTTPRequestHandler):
        def do_GET(self) -> None:
            document = routes.get(self.path)
            if document is None:
                self.send_error(404)
                return

            self.send_response(200)
            self.send_header("Content-Type", document.content_type)
            self.send_header("Content-Length", str(len(document.body)))
            self.end_headers()
            self.wfile.write(document.body)

        def log_message(self, format: str, *args) -> None:
            return

    server = ThreadingHTTPServer(("127.0.0.1", 0), Handler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        host, port = server.server_address
        yield f"http://{host}:{port}"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2.0)


def _demo_documents() -> list[DemoDocument]:
    return [
        DemoDocument(
            path="/article.html",
            content_type="text/html",
            body=(
                b"<html><head><title>Fetch Demo HTML</title></head>"
                b"<body><main><h1>Fetch Demo HTML</h1>"
                b"<p>HTML ingestion reached the vault.</p></main></body></html>"
            ),
            expected_title="Fetch Demo HTML",
            expected_text="HTML ingestion reached the vault.",
        ),
        DemoDocument(
            path="/notes.txt",
            content_type="text/plain",
            body=b"Plain text ingestion reached the vault.",
            expected_title=None,
            expected_text="Plain text ingestion reached the vault.",
        ),
        DemoDocument(
            path="/report.pdf",
            content_type="application/pdf",
            body=_make_pdf_body(
                title="Fetch Demo PDF",
                text="PDF ingestion reached the vault.",
            ),
            expected_title="Fetch Demo PDF",
            expected_text="PDF ingestion reached the vault.",
        ),
        DemoDocument(
            path="/briefing.docx",
            content_type=DOCX_CONTENT_TYPE,
            body=_make_docx_body(
                title="Fetch Demo DOCX",
                paragraph="DOCX ingestion reached the vault.",
            ),
            expected_title="Fetch Demo DOCX",
            expected_text="DOCX ingestion reached the vault.",
        ),
    ]


def _make_docx_body(*, title: str, paragraph: str) -> bytes:
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=1)
    document.add_paragraph(paragraph)
    out = BytesIO()
    document.save(out)
    return out.getvalue()


def _make_pdf_body(*, title: str, text: str) -> bytes:
    """Build a minimal one-page PDF that pypdf can parse and extract."""
    pdf_text = _escape_pdf_string(text)
    pdf_title = _escape_pdf_string(title)
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    stream = f"BT /F1 12 Tf 72 720 Td ({pdf_text}) Tj ET".encode("ascii")
    objects.append(b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream))
    objects.append(f"<< /Title ({pdf_title}) >>".encode("ascii"))

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out.extend(f"{i} 0 obj\n".encode("ascii"))
        out.extend(obj)
        out.extend(b"\nendobj\n")

    xref_offset = len(out)
    out.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    out.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        out.extend(f"{offset:010d} 00000 n \n".encode("ascii"))

    out.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R /Info 6 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(out)


def _escape_pdf_string(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


if __name__ == "__main__":
    raise SystemExit(main())
