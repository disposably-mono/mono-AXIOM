"""
Tests for axiom_fetch.docx_extractor.

DOCX bytes are generated with python-docx in memory so the tests cover
real package parsing without fixture files.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from axiom_fetch.docx_extractor import EmptyExtractionError, extract_docx
from axiom_fetch.extractor import ExtractResult, extract
from docx import Document

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def make_docx(
    paragraphs: list[str] | None = None,
    *,
    title: str | None = "Sample DOCX",
    headings: list[tuple[int, str]] | None = None,
    table_rows: list[list[str]] | None = None,
) -> bytes:
    document = Document()
    if title is not None:
        document.core_properties.title = title

    for level, text in headings or []:
        document.add_heading(text, level=level)

    for text in paragraphs or []:
        document.add_paragraph(text)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
        for row_index, row in enumerate(table_rows):
            for col_index, text in enumerate(row):
                table.cell(row_index, col_index).text = text

    out = BytesIO()
    document.save(out)
    return out.getvalue()


class TestDocxDispatch:
    def test_docx_content_type_routes_to_docx_extractor(self):
        result = extract(make_docx(["Hello from DOCX."]), DOCX_CONTENT_TYPE)

        assert isinstance(result, ExtractResult)
        assert "Hello from DOCX." in result.markdown

    def test_docx_content_type_with_parameters_routes_to_docx(self):
        result = extract(
            make_docx(["DOCX with content type parameters."]),
            f"{DOCX_CONTENT_TYPE}; x=y",
        )

        assert "DOCX with content type parameters." in result.markdown


class TestDocxExtraction:
    def test_extracts_paragraph_text(self):
        result = extract_docx(make_docx(["First paragraph.", "Second paragraph."]))

        assert "First paragraph." in result.markdown
        assert "Second paragraph." in result.markdown

    def test_extracts_core_properties_title(self):
        result = extract_docx(make_docx(["Body text."], title="The DOCX Title"))

        assert result.title == "The DOCX Title"

    def test_missing_core_properties_title_returns_none(self):
        result = extract_docx(make_docx(["Body text."], title=None))

        assert result.title is None

    def test_heading_styles_become_markdown_headings(self):
        result = extract_docx(
            make_docx(["Body text."], headings=[(1, "Chapter One"), (2, "Section")])
        )

        assert "# Chapter One" in result.markdown
        assert "## Section" in result.markdown

    def test_extracts_tables_as_markdown_tables(self):
        result = extract_docx(
            make_docx(
                ["Before table."],
                table_rows=[
                    ["Name", "Role"],
                    ["Ada", "Engineer"],
                ],
            )
        )

        assert "Before table." in result.markdown
        assert "| Name | Role |" in result.markdown
        assert "| --- | --- |" in result.markdown
        assert "| Ada | Engineer |" in result.markdown

    def test_empty_docx_raises_clear_error(self):
        with pytest.raises(EmptyExtractionError, match="no extractable text"):
            extract_docx(make_docx([], title=None))

    def test_invalid_docx_raises_clear_error(self):
        with pytest.raises(ValueError, match="invalid DOCX"):
            extract_docx(b"not a docx")
